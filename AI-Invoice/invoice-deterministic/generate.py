from __future__ import annotations

import os
import shutil
import subprocess
from datetime import datetime, date
from decimal import Decimal, ROUND_HALF_UP
from typing import Dict, Any, Optional, List

from openpyxl import load_workbook
from docx import Document

TWO = Decimal("0.01")


# ---------------- helpers ----------------

def to_date(x: Any) -> date:
    if isinstance(x, datetime):
        return x.date()
    if isinstance(x, date):
        return x
    for fmt in ("%d-%m-%Y", "%Y-%m-%d", "%d/%m/%Y"):
        try:
            return datetime.strptime(str(x), fmt).date()
        except ValueError:
            pass
    raise ValueError(f"Onbekend datumformaat: {x!r}")


def dmy(d: date) -> str:
    return d.strftime("%d-%m-%Y")


def money(x: Decimal) -> str:
    q = x.quantize(TWO, rounding=ROUND_HALF_UP)
    return f"€{q:.2f}".replace(".", ",")


def replace_everywhere(doc: Document, mapping: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                for r in p.runs:
                    r.text = r.text.replace(k, v)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        if k in p.text:
                            for r in p.runs:
                                r.text = r.text.replace(k, v)


# ---------------- excel ----------------

def load_excel(path: str) -> tuple[list[dict], list[dict]]:
    wb = load_workbook(path, data_only=True)
    inv_ws = wb["Facturen"]
    reg_ws = wb["Regels"]

    inv_header = [c.value for c in next(inv_ws.iter_rows(max_row=1))]
    reg_header = [c.value for c in next(reg_ws.iter_rows(max_row=1))]

    invoices, regels = [], []

    for r in inv_ws.iter_rows(min_row=2, values_only=True):
        if r and r[0]:
            invoices.append(dict(zip(inv_header, r)))

    for r in reg_ws.iter_rows(min_row=2, values_only=True):
        if r and r[0]:
            regels.append(dict(zip(reg_header, r)))

    return invoices, regels


# ---------------- word ----------------

def find_lines_table(doc: Document):
    required = {"Datum", "Aantal uren", "Tarief", "BTW"}
    for t in doc.tables:
        header = " ".join(c.text for c in t.rows[0].cells)
        if all(h in header for h in required) and len(t.rows) >= 2:
            return t
    return None


# ---------------- pdf ----------------

def find_soffice() -> Optional[str]:
    p = shutil.which("soffice")
    if p:
        return p
    for c in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]:
        if os.path.exists(c):
            return c
    return None


def docx_to_pdf(paths: List[str], out_dir: str):
    soffice = find_soffice()
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) niet gevonden")

    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        out_dir,
        *paths,
    ]
    subprocess.run(cmd, check=True)


# ---------------- main ----------------

def generate(
    excel_path="facturen.xlsx",
    template_path="template.docx",
    out_dir="out",
):
    os.makedirs(out_dir, exist_ok=True)

    invoices, regels = load_excel(excel_path)

    regels_per_factuur: Dict[str, list] = {}
    for r in regels:
        regels_per_factuur.setdefault(str(r["Factuurnummer"]), []).append(r)

    generated = []

    for inv in invoices:
        inv_no = str(inv["Factuurnummer"])
        lines = regels_per_factuur[inv_no]

        doc = Document(template_path)

        # -------- omschrijving boven tabel --------
        omschrijving = "\n".join(l["Omschrijving"] for l in lines)

        # -------- tabel --------
        table = find_lines_table(doc)
        tmpl = table.rows[1]

        net_total = Decimal("0.00")
        vat_total = Decimal("0.00")

        def fill(row, l):
            nonlocal net_total, vat_total
            qty = Decimal(str(l["Aantal uren"]))
            price = Decimal(str(l["Tarief"]))
            vat_pct = Decimal(str(l["BTW_percentage"]))

            net = qty * price
            vat = net * vat_pct / 100

            net_total += net
            vat_total += vat

            row.cells[0].text = dmy(to_date(l["Datum"]))
            row.cells[1].text = str(qty).replace(".", ",")
            row.cells[2].text = money(price)
            row.cells[3].text = f"{int(vat_pct)}%"
            row.cells[4].text = money(net)

        fill(tmpl, lines[0])
        for l in lines[1:]:
            r = table.add_row()
            fill(r, l)

        gross = net_total + vat_total

        mapping = {
            "{{FACTUURNR}}": inv_no,
            "{{FACTUURDATUM}}": dmy(to_date(inv["Factuurdatum"])),
            "{{VERVALDATUM}}": dmy(to_date(inv["Vervaldatum"])),
            "{{DEBITEURNR}}": str(inv.get("Debiteurnummer", "")),
            "{{KLANT_NAAM}}": inv["Klantnaam"],
            "{{KLANT_ADRES}}": inv["Straat + huisnr"],
            "{{KLANT_POSTCODE_PLAATS}}": inv["Postcode + Plaats"],
            "{{OMSCHRIJVING}}": omschrijving,
            "{{TOTAAL_EXCL}}": money(net_total),
            "{{BTW_TOTAAL}}": money(vat_total),
            "{{TOTAAL_INCL}}": money(gross),
        }

        replace_everywhere(doc, mapping)

        out_docx = os.path.join(out_dir, f"{inv_no}.docx")
        doc.save(out_docx)
        generated.append(out_docx)
        print(f"✅ DOCX: {out_docx}")

    docx_to_pdf(generated, out_dir)
    print("✅ PDF’s gegenereerd")


if __name__ == "__main__":
    generate()
